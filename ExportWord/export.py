#!/usr/bin/env python
# -*- coding:utf-8 -*-
from django.http import JsonResponse, HttpResponse
from django import conf
import datetime, re, os, time, requests, json, uuid
import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

compiles_general = [
    re.compile("^\s*"),
    re.compile("\s*$"),
    re.compile("\t\t"),
    re.compile("\t\t\t*"),
    re.compile("\t \t"),
    re.compile("\t \t \t*"),
]

compiles_enter = [
    re.compile(" \n\n\n*"),
    re.compile(" \n\n "),
    re.compile(" \n\n\n* "),
    re.compile("\n\n*"),
    re.compile("\n \n"),
    re.compile("\n \n \n*"),
]

compiles_two_n = [
    re.compile("\n\n")
]

def WriteFile(ResumeCandidateInfo, filePath):

    fileName = str(uuid.uuid1()) + ".docx"

    newWordFilePath = os.path.join(conf.settings.BASE_DIR, "static/storage", fileName)

    if not os.path.isfile(newWordFilePath):
        doc = Document(filePath)
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        paragraph = doc.add_paragraph()
        paragraph.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run('个人简历')
        run.bold = True
        run.font.name=u'微软雅黑'
        run.font.size=Pt(11)

        # 换行
        doc.add_paragraph()

        BaseInfo = {}
        OtherData = {}

        # 初始化数据
        for k, v in ResumeCandidateInfo.items():
            if k == "BaseInfo":
                # 写入基础数据
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(u"基础信息")
                run.font.name='微软雅黑'
                run.font.size=Pt(11)
                run.bold = True
                for vs in sorted(v.items()):
                    paragraph = doc.add_paragraph()
                    base_key, base_value = vs[1].split("|")[0], vs[1].split("|")[1]
                    run = paragraph.add_run(base_key + ": " + str(base_value) if base_value else "")
                    run.font.name='微软雅黑'
                    run.font.size=Pt(10.5)
            else:
                OtherData[k] = v

        # 换行
        doc.add_paragraph()

        # 写入其他数据
        for kvs in sorted(OtherData.items()):
            long_string_key, long_string_value = kvs[1].split("$|_@")[0], kvs[1].split("$|_@")[1]
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(long_string_key)
            run.bold = True
            run.font.name='微软雅黑'
            run.font.size=Pt(11)

            paragraph = doc.add_paragraph()
            _Data = long_string_value.replace("\r", '').replace("\n\n", '\n')

            # 去除头部换行符
            # _Data = re.sub("^\s*", "", _Data)

            # 去除尾部换行符
            # _Data = re.sub("\s*$", "", _Data)

            # 去除内容不合法换行符
            for i in compiles_general:
                _Data = re.sub(i, "", _Data)
            for i in compiles_enter:
                _Data = re.sub(i, "\n", _Data)
            for i in compiles_two_n:
                _Data = re.sub(i, "\n", _Data)
                
            run = paragraph.add_run(_Data)
            run.font.name='微软雅黑'
            run.font.size=Pt(10.5)

            doc.add_paragraph()

        # 保存文件
        doc.save(newWordFilePath)

    return newWordFilePath

def GenerateWordDocument(request):
    ret = {"status": "seccuss", "status_code": "200"}

    return JsonResponse(ret)